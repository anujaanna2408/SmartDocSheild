import os
import sqlite3
import bcrypt
import logging
import io
import re
from collections import Counter
from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS

try:
    import PyPDF2
    from docx import Document
    import spacy
    import nltk
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize, sent_tokenize
    nlp = spacy.load("en_core_web_sm")
except ImportError:
    nlp = None

# ─── Logging setup ───────────────────────────────────────────────
logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("SmartDocShield")

# ─── App setup ───────────────────────────────────────────────────
app = Flask(__name__, static_folder="static", static_url_path="/static")
CORS(app)

# Check if running on Render
IS_RENDER = os.environ.get('RENDER') == 'true'

# Use the persistent disk mount path on Render, or local 'data' folder
DATA_DIR = "/opt/render/project/src/data" if IS_RENDER else "data"

if not os.path.exists(DATA_DIR):
    os.makedirs(DATA_DIR, exist_ok=True)

DB_NAME = os.path.join(DATA_DIR, "smartdocshield_v2.db")
UPLOAD_FOLDER = os.path.join(DATA_DIR, "uploads")

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def init_db():
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            email TEXT UNIQUE NOT NULL,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            role TEXT DEFAULT 'User',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL,
            filename TEXT NOT NULL,
            custom_name TEXT,
            summary TEXT,
            risk_level TEXT,
            entities_json TEXT,
            pii_json TEXT,
            upload_date DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    conn.commit()
    conn.close()
    log.info("✅ Database v2 initialized securely.")

# Initialize SQLite database on boot
init_db()

# ─── Routes ──────────────────────────────────────────────────────

@app.route("/")
def index():
    return send_from_directory(".", "index.html")

@app.route("/dashboard.html")
def dashboard():
    return send_from_directory("static", "dashboard.html")

@app.route("/<path:path>")
def static_proxy(path):
    if os.path.exists(path):
        return send_from_directory(".", path)
    return jsonify({"error": "File not found"}), 404

@app.route("/api/debug-ls")
def debug_ls():
    import os
    files = os.listdir(".")
    return jsonify({"cwd": os.getcwd(), "files": files})


@app.route("/api/signup", methods=["POST", "OPTIONS"])
def api_signup():
    if request.method == "OPTIONS":
        return jsonify({}), 200

    data = request.json
    if not data:
        return jsonify({"success": False, "error": "No data provided"}), 400

    email = data.get("email", "").strip().lower()
    username = data.get("username", "").strip()
    password = data.get("password", "")
    name = data.get("name", "").strip()

    if not email or "@" not in email:
        return jsonify({"success": False, "error": "Invalid email address."}), 400
    if not username:
        return jsonify({"success": False, "error": "Username is required."}), 400
    if "@" in username:
        return jsonify({"success": False, "error": "Username cannot be an email format."}), 400
    if not name:
        return jsonify({"success": False, "error": "Full name is required."}), 400
    if not password or len(password) < 4:
        return jsonify({"success": False, "error": "Password must be at least 4 characters."}), 400

    log.info(f"📝 DEBUG: New Signup attempt for username: {username}")

    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()

    # Check if user already exists
    cursor.execute("SELECT id FROM users WHERE email = ? OR username = ?", (email, username))
    existing_user = cursor.fetchone()

    if existing_user:
        conn.close()
        log.warning(f"⚠️ Signup Failed: Account with email or username already exists.")
        return jsonify({"success": False, "error": "Email or Username already taken! Please try another or Log In."}), 409

    # Hash the password with bcrypt
    hashed = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

    try:
        cursor.execute(
            "INSERT INTO users (name, email, username, password_hash) VALUES (?, ?, ?, ?)",
            (name, email, username, hashed)
        )
        conn.commit()
        log.info(f"✅ User {username} created successfully inside the SQLite v2 DB.")
        conn.close()
        return jsonify({
            "success": True, 
            "redirect": "index.html",
            "user": {"name": name, "username": username}
        }), 201
    except Exception as e:
        conn.close()
        log.error(f"❌ Database insertion error: {str(e)}")
        return jsonify({"success": False, "error": "Failed to create account. Ensure email/username is unique."}), 500

@app.route("/api/login", methods=["POST", "OPTIONS"])
def api_login():
    if request.method == "OPTIONS":
        return jsonify({}), 200

    data = request.json
    if not data:
        return jsonify({"success": False, "error": "No data provided"}), 400

    username = data.get("username", "").strip()
    password = data.get("password", "")

    if not username:
        return jsonify({"success": False, "error": "Username is required."}), 400
    if not password:
        return jsonify({"success": False, "error": "Password is required."}), 400

    log.info(f"🔍 DEBUG: Login attempt for username: {username}")

    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute("SELECT password_hash, username, name FROM users WHERE username = ?", (username,))
    user = cursor.fetchone()
    conn.close()

    if not user:
        log.warning(f"⚠️ Login Failed: Username '{username}' not found in DB.")
        return jsonify({"success": False, "error": "Account not found. Please click 'Sign Up' first!"}), 404

    stored_hash = user[0]
    db_username = user[1]
    db_name = user[2]
    username = user[1]

    # Compare passwords
    match = bcrypt.checkpw(password.encode('utf-8'), stored_hash.encode('utf-8'))
    log.info(f"🔍 DEBUG: Password match result for {username}: {match}")

    if not match:
        log.warning(f"⚠️  Invalid password submitted for user: {username}")
        return jsonify({"success": False, "error": "Invalid username or password."}), 401

    log.info(f"✅ Credentials verified! User {db_name} ({username}) successfully logged in.")
    return jsonify({
        "success": True, 
        "redirect": "index.html",
        "user": {"name": db_name, "username": db_username}
    }), 200

# ─── Document Processing ─────────────────────────────────────────

import json
import tempfile

# Lazy-load heavy OCR libraries
_easyocr_reader = None

def get_easyocr_reader():
    """Lazy-load EasyOCR reader (heavy init) — singleton pattern."""
    global _easyocr_reader
    if _easyocr_reader is None:
        import easyocr
        _easyocr_reader = easyocr.Reader(['en'], gpu=False, verbose=False)
    return _easyocr_reader


def extract_text_from_image(file_storage, filename):
    """
    Robust image-to-text extraction with multiple fallback strategies.
    Fixes the Windows charmap encoding crash by writing raw bytes directly.
    """
    # Read raw bytes from the uploaded file
    file_bytes = file_storage.read()
    file_storage.seek(0)  # Reset for potential re-use

    if not file_bytes or len(file_bytes) < 100:
        return ""

    ext = os.path.splitext(filename)[1].lower()
    if not ext:
        ext = ".png"

    # Write to a temp file safely (binary mode, no encoding issues)
    tmp_path = os.path.join(tempfile.gettempdir(), f"sds_ocr_{os.getpid()}{ext}")
    try:
        with open(tmp_path, "wb") as f:
            f.write(file_bytes)
    except Exception as e:
        log.error(f"Failed to write temp file: {e}")
        return ""

    extracted_text = ""

    # ── Strategy 1: EasyOCR (best for complex layouts, forms, IDs) ──
    try:
        reader = get_easyocr_reader()
        results = reader.readtext(tmp_path, detail=1, paragraph=False)
        # Sort results by vertical position (top-to-bottom, left-to-right)
        results.sort(key=lambda r: (r[0][0][1], r[0][0][0]))
        
        lines = []
        current_line = []
        last_y = None
        y_threshold = 15  # pixels tolerance for same-line grouping
        
        for (bbox, text, confidence) in results:
            if confidence < 0.15:  # Skip very low confidence noise
                continue
            top_y = bbox[0][1]
            if last_y is not None and abs(top_y - last_y) > y_threshold:
                # New line detected
                lines.append(" ".join(current_line))
                current_line = []
            current_line.append(text.strip())
            last_y = top_y
        
        if current_line:
            lines.append(" ".join(current_line))
        
        extracted_text = "\n".join(lines)
        log.info(f"EasyOCR extracted {len(extracted_text)} chars from {filename}")
    except Exception as e:
        log.warning(f"EasyOCR failed for {filename}: {e}")
        extracted_text = ""

    # ── Strategy 2: Tesseract fallback (if EasyOCR gave nothing) ──
    if not extracted_text.strip():
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(tmp_path)
            # Preprocess for better OCR: convert to grayscale
            img = img.convert("L")
            extracted_text = pytesseract.image_to_string(img, lang="eng")
            log.info(f"Tesseract extracted {len(extracted_text)} chars from {filename}")
        except Exception as e:
            log.warning(f"Tesseract fallback also failed for {filename}: {e}")

    # Cleanup temp file
    try:
        os.unlink(tmp_path)
    except:
        pass

    return extracted_text.strip()


def detect_pii(text):
    """
    Comprehensive PII detection with patterns for Indian and international documents.
    Returns list of PII matches and the risk level.
    """
    pii_matches = []
    seen_values = set()

    # Define all PII patterns
    patterns = [
        # Indian PII
        ("PAN Card",      r"\b[A-Z]{5}\d{4}[A-Z]\b"),
        ("Aadhaar",       r"\b\d{4}\s?\d{4}\s?\d{4}\b"),
        # International
        ("SSN",           r"\b\d{3}-\d{2}-\d{4}\b"),
        ("Phone",         r"\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b"),
        ("Email",         r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}"),
        ("Credit Card",   r"\b(?:\d{4}[-\s]?){3}\d{4}\b"),
        ("Passport",      r"\b[A-Z]{1,2}\d{7,8}\b"),
        ("Date of Birth", r"\b(?:DOB|D\.O\.B|Date of Birth|Birth Date)[:\s]*\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{2,4}\b"),
        ("Date of Birth", r"\b\d{1,2}[/\-\.]\d{1,2}[/\-\.]\d{4}\b"),
    ]

    for pii_type, pattern in patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            value = match.group().strip()
            # Skip if already seen or too short
            if value in seen_values or len(value) < 3:
                continue
            # For Aadhaar, verify it's exactly 12 digits
            if pii_type == "Aadhaar":
                digits = re.sub(r"\s", "", value)
                if len(digits) != 12 or not digits.isdigit():
                    continue
            # For PAN, must be exact 10 chars
            if pii_type == "PAN Card" and len(value) != 10:
                continue
            seen_values.add(value)
            pii_matches.append({"type": pii_type, "value": value})

    # Calculate risk level
    if len(pii_matches) > 5:
        risk_level = "Critical"
    elif len(pii_matches) > 3:
        risk_level = "High"
    elif len(pii_matches) > 0:
        risk_level = "Medium"
    else:
        risk_level = "Low"

    return pii_matches, risk_level


def mask_pii_in_text(text, pii_matches):
    """Replace detected PII with masked placeholders."""
    masked = text
    for pii in pii_matches:
        val = pii['value']
        ptype = pii['type']
        if ptype == 'SSN':
            masked = masked.replace(val, "XXX-XX-XXXX")
        elif ptype == 'Phone':
            masked = masked.replace(val, "(XXX) XXX-XXXX")
        elif ptype == 'Email':
            masked = masked.replace(val, "[REDACTED@EMAIL.COM]")
        elif ptype == 'PAN Card':
            masked = masked.replace(val, "XXXXX" + val[5:9] + "X")
        elif ptype == 'Aadhaar':
            masked = masked.replace(val, "XXXX XXXX " + val[-4:])
        elif ptype == 'Credit Card':
            masked = masked.replace(val, "XXXX-XXXX-XXXX-" + val[-4:])
        elif ptype == 'Passport':
            masked = masked.replace(val, val[0] + "XXXXXXX")
        elif ptype == 'Date of Birth':
            masked = masked.replace(val, "[DOB REDACTED]")
        else:
            masked = masked.replace(val, "[REDACTED]")
    return masked


def extract_entities_structured(text):
    """
    Extract named entities using SpaCy and group them by category.
    Returns both a flat list and a grouped dictionary for structured display.
    """
    entities_flat = []
    entities_grouped = {}

    if nlp is None:
        return entities_flat, entities_grouped

    try:
        # SpaCy has a max length; truncate very long texts
        analysis_text = text[:100000] if len(text) > 100000 else text
        doc_nlp = nlp(analysis_text)
        seen_texts = set()

        # Map SpaCy labels to friendly categories
        label_map = {
            "PERSON": "Person / Name",
            "ORG": "Organization",
            "GPE": "Location (City/Country)",
            "LOC": "Location (Geographic)",
            "FAC": "Facility / Address",
            "DATE": "Date",
            "MONEY": "Monetary Amount",
            "CARDINAL": "Number",
            "ORDINAL": "Ordinal",
            "PERCENT": "Percentage",
            "TIME": "Time",
            "QUANTITY": "Quantity",
            "PRODUCT": "Product",
            "EVENT": "Event",
            "WORK_OF_ART": "Document / Work",
            "LAW": "Legal Reference",
            "LANGUAGE": "Language",
            "NORP": "Group / Nationality",
        }

        for ent in doc_nlp.ents:
            clean_text = ent.text.strip().replace('\n', ' ')
            if not clean_text or clean_text in seen_texts or len(clean_text) < 2:
                continue
            # Skip entities that are just numbers or single chars
            if clean_text.isdigit() and len(clean_text) < 3:
                continue
            seen_texts.add(clean_text)
            label = label_map.get(ent.label_, ent.label_)
            entities_flat.append({"text": clean_text, "label": label})

            if label not in entities_grouped:
                entities_grouped[label] = []
            entities_grouped[label].append(clean_text)

        # Add key phrases from noun chunks (limited)
        phrase_count = 0
        for chunk in doc_nlp.noun_chunks:
            phrase = chunk.text.strip().replace('\n', ' ')
            if len(phrase.split()) >= 2 and phrase not in seen_texts and phrase_count < 8:
                seen_texts.add(phrase)
                entities_flat.append({"text": phrase, "label": "Key Phrase"})
                if "Key Phrase" not in entities_grouped:
                    entities_grouped["Key Phrase"] = []
                entities_grouped["Key Phrase"].append(phrase)
                phrase_count += 1

    except Exception as e:
        log.error(f"SpaCy entity extraction error: {e}")

    return entities_flat, entities_grouped


def generate_summary(text):
    """
    Generate an extractive summary using NLTK.
    Picks the most informative sentences from the text.
    """
    if not text or len(text.strip()) < 20:
        return "Document content is too short for meaningful summarization."

    try:
        sentences = sent_tokenize(text)
        if len(sentences) <= 2:
            return text.strip()

        stop_words = set(stopwords.words("english"))
        words = word_tokenize(text.lower())
        freq = Counter(w for w in words if w.isalnum() and w not in stop_words and len(w) > 2)

        if not freq:
            return sentences[0] if sentences else text.strip()

        # Score each sentence
        scores = {}
        for i, sentence in enumerate(sentences):
            sentence_words = word_tokenize(sentence.lower())
            score = sum(freq.get(w, 0) for w in sentence_words if w.isalnum())
            # Normalize by sentence length to avoid bias toward long sentences
            word_count = max(len(sentence_words), 1)
            scores[i] = score / word_count

        # Pick top 3-5 sentences depending on document length
        num_summary_sentences = min(max(3, len(sentences) // 4), 5)
        top_indices = sorted(scores.keys(), key=lambda k: scores[k], reverse=True)[:num_summary_sentences]
        top_indices.sort()  # Maintain original order

        summary_sentences = [sentences[i].replace('\n', ' ').strip() for i in top_indices]
        return " ".join(summary_sentences)
    except Exception as e:
        log.error(f"NLTK summarization error: {e}")
        return text[:500].strip() + "..." if len(text) > 500 else text.strip()


def detect_document_type(text, filename):
    """Attempt to classify the document type based on content and filename."""
    text_lower = text.lower()
    fname_lower = filename.lower()

    if any(kw in text_lower for kw in ["pan", "permanent account number", "income tax"]):
        return "PAN Card"
    if any(kw in text_lower for kw in ["aadhaar", "unique identification", "uidai"]):
        return "Aadhaar Card"
    if any(kw in text_lower for kw in ["passport", "republic of india", "nationality"]):
        return "Passport"
    if any(kw in text_lower for kw in ["invoice", "bill to", "total amount", "subtotal"]):
        return "Invoice / Bill"
    if any(kw in text_lower for kw in ["resume", "curriculum vitae", "work experience", "education"]):
        return "Resume / CV"
    if any(kw in text_lower for kw in ["contract", "agreement", "terms and conditions", "hereby"]):
        return "Legal Contract"
    if any(kw in text_lower for kw in ["bank statement", "account number", "transaction"]):
        return "Bank Statement"
    if any(kw in text_lower for kw in ["medical", "diagnosis", "patient", "prescription"]):
        return "Medical Document"
    if any(kw in text_lower for kw in ["certificate", "certify", "awarded"]):
        return "Certificate"
    if any(kw in fname_lower for kw in ["report", "analysis"]):
        return "Report"
    return "General Document"


@app.route("/api/upload", methods=["POST", "OPTIONS"])
def api_upload():
    if request.method == "OPTIONS":
        return jsonify({}), 200

    username = request.form.get('username', 'guest')
    custom_name = request.form.get('custom_name', '')

    files = request.files.getlist('files')
    if not files:
        if 'file' in request.files:
            files = [request.files['file']]
        else:
            return jsonify({"success": False, "error": "No file payloads detected!"}), 400

    processing_results = []

    for file in files:
        original_filename = file.filename
        filename = original_filename.lower() if original_filename else ""
        if not filename:
            continue

        log.info(f"Processing document: {original_filename}")
        
        import uuid
        safe_name = f"{uuid.uuid4().hex[:8]}_{original_filename}"
        file_path = os.path.join(UPLOAD_FOLDER, safe_name)
        file.save(file_path)

        text = ""

        # ════════════════════════════════════════════════════════════
        # STEP 1: Text Extraction
        # ════════════════════════════════════════════════════════════
        try:
            with open(file_path, "rb") as saved_file:
                if filename.endswith(".pdf"):
                    reader = PyPDF2.PdfReader(saved_file)
                    for page in reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + "\n"
                elif filename.endswith(".docx"):
                    doc = Document(saved_file)
                    for p in doc.paragraphs:
                        text += p.text + "\n"
                    # Also extract from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                            if row_text:
                                text += row_text + "\n"
                elif filename.endswith(".txt"):
                    text = saved_file.read().decode('utf-8', errors='ignore')
                elif filename.endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp")):
                    text = extract_text_from_image(saved_file, filename)
                else:
                    processing_results.append({
                        "filename": original_filename,
                        "error": f"Unsupported file format: {os.path.splitext(filename)[1]}"
                    })
                    continue
        except Exception as e:
            log.error(f"Extraction failed for {original_filename}: {e}")
            processing_results.append({
                "filename": original_filename,
                "error": f"Failed to parse document: {str(e)}"
            })
            continue

        if not text.strip():
            processing_results.append({
                "filename": original_filename,
                "error": "No readable text could be extracted from this document. The image may be too blurry, corrupt, or contain no text."
            })
            continue

        # Clean up extracted text (remove excessive whitespace)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = text.strip()

        log.info(f"Extracted {len(text)} characters from {original_filename}. Running analysis...")

        # ════════════════════════════════════════════════════════════
        # STEP 2: Document Type Detection
        # ════════════════════════════════════════════════════════════
        doc_type = detect_document_type(text, filename)
        log.info(f"Document classified as: {doc_type}")

        # ════════════════════════════════════════════════════════════
        # STEP 3: PII Detection & Masking
        # ════════════════════════════════════════════════════════════
        pii_matches, risk_level = detect_pii(text)
        masked_text = mask_pii_in_text(text, pii_matches)

        log.info(f"PII scan: {len(pii_matches)} items found. Risk: {risk_level}")

        # ════════════════════════════════════════════════════════════
        # STEP 4: Entity Extraction (SpaCy NLP)
        # ════════════════════════════════════════════════════════════
        entities_flat, entities_grouped = extract_entities_structured(masked_text)

        # ════════════════════════════════════════════════════════════
        # STEP 5: Summarization (NLTK Extractive)
        # ════════════════════════════════════════════════════════════
        summary = generate_summary(masked_text)

        # ════════════════════════════════════════════════════════════
        # STEP 6: Persist to Database
        # ════════════════════════════════════════════════════════════
        if username != 'guest':
            try:
                conn = sqlite3.connect(DB_NAME)
                c = conn.cursor()
                c.execute('''INSERT INTO documents
                            (username, filename, custom_name, summary, risk_level, entities_json, pii_json)
                            VALUES (?, ?, ?, ?, ?, ?, ?)''',
                          (username, filename, custom_name, summary, risk_level,
                           json.dumps(entities_flat), json.dumps(pii_matches)))
                conn.commit()
                conn.close()
            except Exception as e:
                log.error(f"Database error: {e}")

        log.info(f"Analysis complete for {original_filename}! Risk: {risk_level}")

        processing_results.append({
            "success": True,
            "filename": original_filename,
            "custom_name": custom_name,
            "doc_type": doc_type,
            "text_length": len(text),
            "extracted_text": masked_text,
            "summary": summary,
            "entities": entities_flat,
            "entities_grouped": entities_grouped,
            "pii": pii_matches,
            "risk_level": risk_level,
        })

    return jsonify({
        "success": True,
        "results": processing_results
    }), 200

# ─── Document Archives History ──────────────────────────────────

@app.route("/api/documents", methods=["GET"])
def api_documents():
    username = request.args.get("username")
    query = request.args.get("q", "").strip()
    if not username:
        return jsonify({"success": False, "error": "Username is required to access archives."}), 400

    try:
        conn = sqlite3.connect(DB_NAME)
        c = conn.cursor()
        
        if query:
            search_pattern = f"%{query}%"
            c.execute('''SELECT id, filename, custom_name, summary, risk_level, upload_date 
                         FROM documents 
                         WHERE username=? AND (filename LIKE ? OR custom_name LIKE ?)
                         ORDER BY id DESC''', (username, search_pattern, search_pattern))
        else:
            c.execute('''SELECT id, filename, custom_name, summary, risk_level, upload_date 
                         FROM documents WHERE username=? ORDER BY id DESC''', (username,))
        
        rows = c.fetchall()
        conn.close()
        
        docs = []
        for r in rows:
            docs.append({
                "id": r[0],
                "filename": r[1],
                "custom_name": r[2] if r[2] else "",
                "summary": r[3],
                "risk_level": r[4],
                "upload_date": r[5]
            })
        return jsonify({"success": True, "documents": docs}), 200
    except Exception as e:
        log.error(f"❌ Failed to fetch documents: {str(e)}")
        return jsonify({"success": False, "error": str(e)}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True, use_reloader=False)
